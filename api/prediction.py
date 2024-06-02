import docx
import glob
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_UNDERLINE

import os
from shutil import move

import pandas as pd
import matplotlib.pyplot as plt
import numpy as np


def create_docx(res):
    df = pd.DataFrame()

    # Initialise the Word document
    doc = docx.Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(14)

    doc.add_heading('Результаты прогноза', 0)

    p = doc.add_paragraph('ФИО: ' )
    run = p.add_run('Тим С.Э.\n')
    run.font.underline = WD_UNDERLINE.SINGLE
    run.font.bold = True

    p.add_run('Дата рождения: ')
    run = p.add_run('20.03.1953')
    run.font.underline = WD_UNDERLINE.SINGLE
    run.font.bold = True

    p.add_run('    Возраст: ')
    run = p.add_run('59\n')
    run.font.underline = WD_UNDERLINE.SINGLE
    run.font.bold = True

    p.add_run('Вес: ') 
    run = p.add_run('70')
    run.font.underline = WD_UNDERLINE.SINGLE
    run.font.bold = True

    p.add_run('    Рост: ')
    run = p.add_run('187')
    run.font.underline = WD_UNDERLINE.SINGLE
    run.font.bold = True

    p.add_run('    Адрес: ') 
    run = p.add_run('Аякс\n')
    run.font.underline = WD_UNDERLINE.SINGLE
    run.font.bold = True

    p.add_run('Дата обследования: ')
    run = p.add_run('01.06.2024\n')
    run.font.underline = WD_UNDERLINE.SINGLE
    run.font.bold = True

    p.add_run('Длительность наблюдения: ')
    run = p.add_run('6:21:06\n')
    run.font.underline = WD_UNDERLINE.SINGLE 
    run.font.bold = True

    p.add_run(res)
    p.add_run("\n")

    p.add_run('Вероятность наличия ')
    p.add_run('синдрома сонного апноэ: ').italic = True
    p.add_run('50%').bold = True

    for i in range(18):
        p.add_run('\n')

    for i in range(53):
        p.add_run(' ')

    p.add_run('1.06.2024     ') 
    run = p.add_run('    Врач: _______________')

    # Initialise the table
    t = doc.add_table(rows=1, cols=df.shape[1])
    # Add the column headings
    for j in range(df.shape[1]):
        cell = df.columns[j]
        p = t.cell(0, j).add_paragraph('')
        p.add_run(str(cell)).bold = True
    # Add the body of the data frame
    for i in range(df.shape[0]):
        row = t.add_row()
        for j in range(df.shape[1]):
            cell = df.iat[i, j]
            row.cells[j].text = str(cell)

    # Save the Word doc
    doc.save('result.docx')



